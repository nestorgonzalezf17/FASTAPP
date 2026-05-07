from pathlib import Path
import docx, json, re

import os

def limpiar():
    # 'nt' es para Windows, 'posix' para Linux/macOS
    os.system('cls' if os.name == 'nt' else 'clear')

error_files = {}
dir_errors = "./Results/Error_files.json"
dir_json = "./Results/docx_files.json"
dir_antijob = "./Results/AntiJob.json"
dir_missing_docs = "./Results/MissingDocs.json"
dir_sql = "./Results/InsertFunctions.sql"
dir_resultados = "./SchemsData/Resultados.json"
dir_missing_txt = "./Results/MissingDocs_Summary.txt"

def normalize_text(text):
    """Limpia el texto eliminando espacios múltiples, saltos de línea y convirtiendo a formato estándar."""
    if not text:
        return ""
    # Reemplazar múltiples espacios/saltos por uno solo y quitar espacios en los extremos
    clean = " ".join(text.split())
    return clean

insert_line_re = re.compile(
    r"VALUES\s*\(\s*(?P<id_cargo>\d+)\s*,\s*(?P<id_competencia>\d+)\s*,\s*N'(?P<texto>(?:''|[^'])*?)'\s*,\s*(?P<orden>\d+)\s*,\s*(?P<activo>\d+)\s*,\s*GETDATE\(\)\s*,\s*(?P<id_empresa>\d+)\s*\);",
    re.IGNORECASE
)

def parse_sql_insert_line(line):
    match = insert_line_re.search(line)
    if not match:
        return None
    data = match.groupdict()
    data['id_cargo'] = int(data['id_cargo'])
    data['id_competencia'] = int(data['id_competencia'])
    data['orden'] = int(data['orden'])
    data['activo'] = int(data['activo'])
    data['id_empresa'] = int(data['id_empresa'])
    data['texto'] = data['texto'].replace("''", "'")
    return data


def escape_sql_text(text):
    return text.replace("'", "''")


def conjunction_dbs(data):
    if not os.path.exists(dir_resultados):
        print(f"Error: No se encontró el archivo de referencia {dir_resultados}")
        return

    with open(dir_resultados, 'r', encoding='utf-8-sig') as f:
        resultados = json.load(f)
    
    # 1. Obtener set de cargos únicos en la DB (por nombre)
    DB_jobs_set = set()
    for r in resultados:
        cargo_name = str(r.get('Nombre', '')).strip().upper()
        if cargo_name:
            DB_jobs_set.add(cargo_name)
    
    AntiJob = {}
    Errors = {}
    MissingDocs = {}
    
    # 2. Identificar Nuevos Cargos (en docx pero no en DB por NOMBRE)
    docx_jobs_found = set() # Todos los cargos encontrados en DOCX
    for enterprise_key, enterprise_data in data.items():
        if isinstance(enterprise_data, dict):
            for job_docx, functions in enterprise_data.items():
                job_upper = job_docx.strip().upper()
                docx_jobs_found.add(job_upper)
                
                if job_upper not in DB_jobs_set:
                    AntiJob.setdefault(enterprise_key, {}).setdefault("Nuevos Cargos", {})[job_docx] = functions
        else:
            Errors[enterprise_key] = enterprise_data

    # 3. Identificar Cargos Faltantes (en DB pero no en docx por NOMBRE)
    for r in resultados:
        cargo_db = str(r.get('Nombre', '')).strip().upper()
        emp_db = str(r.get('NombreEmpresa', 'DESCONOCIDA')).strip().upper()
        area_db = str(r.get('NombreArea', 'SIN AREA')).strip().upper()
        
        if cargo_db not in docx_jobs_found:
            MissingDocs.setdefault(emp_db, {}).setdefault(area_db, []).append(cargo_db)


    # Guardar archivos
    if Errors:
        with open(dir_errors, 'w', encoding='utf-8') as f:
            json.dump(Errors, f, ensure_ascii=False, indent=2)

    with open(dir_antijob, 'w', encoding='utf-8') as f:
        json.dump(AntiJob, f, ensure_ascii=False, indent=2)
        
    with open(dir_missing_docs, 'w', encoding='utf-8') as f:
        json.dump(MissingDocs, f, ensure_ascii=False, indent=2)
    
    print(f"\n--- Resumen del Proceso ---")
    print(f"Empresas con cargos nuevos detectados (en DOCX pero no en DB): {len(AntiJob)}")
    print(f"Empresas con cargos faltantes de documentación (en DB pero no en DOCX): {len(MissingDocs)}")
    
    # Detalle en consola y generación de TXT
    missing_summary_lines = []
    
    if MissingDocs:
        print("\n--- DETALLE DE CARGOS FALTANTES (En DB pero sin archivo DOCX) ---")
        for emp, areas in MissingDocs.items():
            for area, jobs in areas.items():
                for job in jobs:
                    line = f"{emp} | {job}"
                    missing_summary_lines.append(line)
                    print(f" [FALTANTE] {line}")
    
    with open(dir_missing_txt, 'w', encoding='utf-8') as f:
        f.write("\n".join(missing_summary_lines))
        
    if AntiJob:
        print("\n--- DETALLE DE CARGOS NUEVOS (En archivo DOCX pero no en DB) ---")
        for emp, content in AntiJob.items():
            if isinstance(content, dict) and "Nuevos Cargos" in content:
                for job in content["Nuevos Cargos"].keys():
                    print(f" [NUEVO] {emp} | {job}")

    print(f"\nReportes guardados en la carpeta /Results")
    print(f"Resumen de faltantes generado en: {dir_missing_txt}")


def validate_no_overlap():
    """Valida que no haya cargos que existan en docx_files.json y MissingDocs.json al mismo tiempo."""
    if not os.path.exists(dir_json) or not os.path.exists(dir_missing_docs):
        return

    with open(dir_json, 'r', encoding='utf-8') as f:
        docx_data = json.load(f)
    with open(dir_missing_docs, 'r', encoding='utf-8') as f:
        missing_data = json.load(f)

    # Crear set de todos los cargos en docx_data
    docx_jobs = set()
    for emp, jobs in docx_data.items():
        for job in jobs.keys():
            docx_jobs.add(job.strip().upper())

    overlaps = []
    
    for emp, areas in missing_data.items():
        for area, jobs in areas.items():
            for job in jobs:
                if job.strip().upper() in docx_jobs:
                    overlaps.append(f"[{emp}] {job} (Área: {area})")

    if overlaps:
        print(f"\n[ALERTA] Se encontraron {len(overlaps)} solapamientos críticos:")
        for overlap in overlaps:
            print(f" - {overlap}")
        
        # Guardar reporte de solapamiento
        with open("./Results/Overlap_Validation.json", 'w', encoding='utf-8') as f:
            json.dump({"overlaps": overlaps, "count": len(overlaps)}, f, ensure_ascii=False, indent=2)
    else:
        print("\n[ÉXITO] Validación de consistencia completada: No hay solapamientos entre DOCX y MissingDocs.")


def generate_sql_script(data):
    if not os.path.exists(dir_resultados):
        print(f"Error: No se encontró el archivo de referencia {dir_resultados}")
        return

    with open(dir_resultados, 'r', encoding='utf-8-sig') as f:
        resultados = json.load(f)

    lookup = {}
    for r in resultados:
        # Usamos solo el nombre del cargo para el matching global
        cargo = str(r.get('Nombre', '')).strip().upper()
        if cargo:
            # Guardamos una lista de tuplas (IdEmpresa, IdCargo, NombreEmpresa) por si hay duplicados
            lookup.setdefault(cargo, []).append({
                'id_empresa': r['IdEmpresa'],
                'id_cargo': r['IdCargo'],
                'empresa': r.get('NombreEmpresa', 'DESCONOCIDA')
            })

    sql_statements = [
        "-- SCRIPT DE INSERCIÓN DE FUNCIONES",
        "-- Generado automáticamente",
        "SET ANSI_NULLS ON",
        "GO",
        "SET QUOTED_IDENTIFIER ON",
        "GO",
        ""
    ]

    count_cargos = 0
    count_items = 0
    
    # 1. Consolidar todas las funciones por cargo de forma global
    # Esto evita duplicados si el mismo cargo aparece en carpetas de distintas empresas
    global_docx_data = {} # job_key -> map(normalized_text -> original_text)
    for enterprise_name, jobs in data.items():
        if not isinstance(jobs, dict):
            continue
        for job_title, functions in jobs.items():
            job_key = job_title.strip().upper()
            if job_key not in global_docx_data:
                global_docx_data[job_key] = {}
            for f in functions:
                normalized_f = normalize_text(f)
                if normalized_f:
                    # Usamos la versión normalizada como clave para evitar duplicados por espacios/formato
                    # Mantenemos el texto original (con su capitalización) si es la primera vez que lo vemos
                    norm_key = normalized_f.upper()
                    if norm_key not in global_docx_data[job_key]:
                        global_docx_data[job_key][norm_key] = normalized_f

    # 2. Generar SQL usando la data consolidada
    for job_key, functions_map in global_docx_data.items():
        if job_key in lookup:
            # Para cada cargo encontrado en la DB con ese nombre, generar el script
            for record in lookup[job_key]:
                id_empresa = record['id_empresa']
                id_cargo = record['id_cargo']
                emp_db_name = record['empresa']
                
                sql_statements.append(f"-- EMPRESA DB: {emp_db_name} | CARGO: {job_key}")
                
                # Convertir el mapa a lista de textos originales
                functions_list = sorted(list(functions_map.values()))
                
                for i, function_text in enumerate(functions_list, 1):
                    clean_text = escape_sql_text(function_text)
                    sql = (f"INSERT INTO [EDD].[ItemsEvaluacionCargo] "
                           f"([IdCargo], [IdCompetencia], [TextoItem], [Orden], [Activo], [FechaCreacion], [IdEmpresa]) "
                           f"VALUES ({id_cargo}, 1, N'{clean_text}', {i}, 1, GETDATE(), {id_empresa});")
                    sql_statements.append(sql)
                    count_items += 1
                
                sql_statements.append("") # Separador
            count_cargos += 1
        else:
            # print(f"Aviso: No se encontró ID para el cargo {job_key} en la DB")
            pass

    with open(dir_sql, 'w', encoding='utf-8') as f:
        f.write("\n".join(sql_statements))
    
    print(f"Script SQL generado: {dir_sql}")
    print(f"Se generaron {count_items} inserciones para {count_cargos} cargos.")


def corregir_duplicados_orden():
    """Corrige los valores de Orden en el SQL generado para evitar duplicados en (IdCargo, IdCompetencia, Orden, Activo, IdEmpresa)."""
    if not os.path.exists(dir_sql):
        print(f"Error: No se encontró el archivo {dir_sql}")
        return

    with open(dir_sql, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    grupos = {}

    for i, line in enumerate(lines):
        parsed = parse_sql_insert_line(line)
        if not parsed:
            continue

        key = (parsed['id_cargo'], parsed['id_empresa'])
        grupos.setdefault(key, []).append((parsed['orden'], i))

    for items in grupos.values():
        items.sort(key=lambda x: x[0])
        for new_orden, (_, line_idx) in enumerate(items, 1):
            parsed = parse_sql_insert_line(lines[line_idx])
            if not parsed:
                continue

            texto_sql = escape_sql_text(parsed['texto'])
            lines[line_idx] = (
                f"INSERT INTO [EDD].[ItemsEvaluacionCargo] "
                f"([IdCargo], [IdCompetencia], [TextoItem], [Orden], [Activo], [FechaCreacion], [IdEmpresa]) "
                f"VALUES ({parsed['id_cargo']}, {parsed['id_competencia']}, N'{texto_sql}', {new_orden}, {parsed['activo']}, GETDATE(), {parsed['id_empresa']});\n"
            )

    with open(dir_sql, 'w', encoding='utf-8') as f:
        f.writelines(lines)

    print("Duplicados de Orden corregidos en el script SQL.")


def cargar_mapeo_cargos():
    if not os.path.exists(dir_resultados):
        return {}

    with open(dir_resultados, 'r', encoding='utf-8-sig') as f:
        resultados = json.load(f)

    mapping = {}
    for r in resultados:
        key = (r.get('IdEmpresa'), r.get('IdCargo'))
        mapping[key] = {
            'empresa': r.get('NombreEmpresa', 'DESCONOCIDA'),
            'area': r.get('NombreArea', 'SIN AREA'),
            'cargo': r.get('Nombre', f'Cargo {r.get("IdCargo", "?")}')
        }
    return mapping


def generar_reporte_items():
    """Genera un TXT con la estructura Empresa > Area > Cargo > item - TextoItem y valida los items."""
    if not os.path.exists(dir_sql):
        print(f"Error: No se encontró el archivo {dir_sql}")
        return

    mapping = cargar_mapeo_cargos()

    estructura = {}
    problemas = []

    with open(dir_sql, 'r', encoding='utf-8') as f:
        for line in f:
            parsed = parse_sql_insert_line(line)
            if not parsed:
                continue

            key = (parsed['id_empresa'], parsed['id_cargo'])
            meta = mapping.get(key, {
                'empresa': f'Empresa {parsed["id_empresa"]}',
                'area': 'SIN AREA',
                'cargo': f'Cargo {parsed["id_cargo"]}'
            })

            empresa = meta['empresa']
            area = meta['area']
            cargo = meta['cargo']

            estructura.setdefault(empresa, {}).setdefault(area, {}).setdefault(cargo, []).append(
                (parsed['orden'], parsed['texto'])
            )

    report_lines = []
    for empresa in sorted(estructura):
        report_lines.append(f"Empresa: {empresa}")
        for area in sorted(estructura[empresa]):
            report_lines.append(f"    Area: {area}")
            for cargo in sorted(estructura[empresa][area]):
                items = sorted(estructura[empresa][area][cargo], key=lambda x: x[0])
                ordenes = [orden for orden, _ in items]
                referencia = list(range(1, len(items) + 1))
                duplicados = sorted({orden for orden in ordenes if ordenes.count(orden) > 1})

                if ordenes != referencia or duplicados:
                    report_lines.append(f"        Cargo: {cargo}  [VALIDACIÓN FALLIDA]")
                    if duplicados:
                        report_lines.append(f"            Items duplicados: {', '.join(map(str, duplicados))}")
                    if ordenes != referencia:
                        report_lines.append(f"            Secuencia esperada: {', '.join(map(str, referencia))}")
                        report_lines.append(f"            Secuencia actual:   {', '.join(map(str, ordenes))}")
                    problemas.append((empresa, area, cargo, ordenes))
                else:
                    report_lines.append(f"        Cargo: {cargo}")

                for orden, texto in items:
                    report_lines.append(f"            {orden} - {texto}")
                report_lines.append("")

    if problemas:
        report_lines.append("VALIDACIÓN: Se encontraron cargos con items no consecutivos o duplicados.")
        report_lines.append(f"Cargos con problemas: {len(problemas)}")
    else:
        report_lines.append("VALIDACIÓN: Todos los items están en orden y sin duplicados.")

    report_path = "./Results/ItemsReport.txt"
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(report_lines))

    print(f"Reporte de items generado: {report_path}")



def list_funtions(text):
    lines = text.splitlines()
    functions = []
    seen = set()
    for line in lines:
        clean = normalize_text(line)
        if clean:
            norm_key = clean.upper()
            if norm_key not in seen:
                functions.append(clean)
                seen.add(norm_key)
    return functions

def search_funtions_job(doc, text):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.tables:
                    for subtable in cell.tables:
                        i = 0
                        for subrow in subtable.rows:
                            j = 0
                            for subcell in subrow.cells: 
                                if "FUNCIONES DEL CARGO" in subcell.text:
                                    pass

                                if i == 1 and j == 0:
                                    text[list(text.keys())[0]] = list_funtions(subcell.text)
                                j += 1
                            i += 1

def search_job_title(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                
                if "DENOMINACIÓN DEL CARGO:"  in cell.text:
                    cleaned_title = cell.text.removeprefix("DENOMINACIÓN DEL CARGO:").strip()
                    return {cleaned_title: ""}  
    return "No se encontró el título del cargo."

def extract_text_from_docx(file):
    try:
        doc = docx.Document(file)
        text = search_job_title(doc)
        search_funtions_job(doc, text)
        #text.update(search_funtions_job(doc, text))
        return text
    except Exception as e:
        return f"Error extracting text from {file}: {e}"

def is_docx(file):
    return file.suffix == ".docx"

def list_files(directory, docxs=None):
    if docxs is None:
        docxs = {}
    ErrorValues = ["Error extracting text from", "Not a .docx file"]
    for element in directory.iterdir():
        if element.is_file():
            if is_docx(element):
                text = extract_text_from_docx(element)
                val = False
                for error in ErrorValues:
                    if error in text:
                        error_files[str(element)] = text
                        val = True
                if val == False:
                    #print(f"@ Text from {element}: {text}")
                    PahtParts = str(element).split("\\")
                    Enterprice = "DESCONOCIDA"
                    for part in PahtParts:
                        if "PERFILES" in part.upper():
                            raw_name = part.replace("PERFILES", "").strip().upper()
                            # Normalización de nombres de empresa
                            if "GEZPOMOTOR" in raw_name:
                                Enterprice = "GEZPMOTOR"
                            else:
                                Enterprice = raw_name
                            break
                    
                    job = (list(text.keys()))[0]
                    
                    if Enterprice not in docxs:
                        docxs[Enterprice] = {}
                    
                    if job not in docxs[Enterprice]:
                        docxs[Enterprice][job] = text[job]
                    else:
                        # Si el cargo ya existe, añadir funciones que no estén repetidas (normalizado)
                        seen_functions = {normalize_text(f).upper() for f in docxs[Enterprice][job]}
                        for function in text[job]:
                            norm_f = normalize_text(function)
                            if norm_f.upper() not in seen_functions:
                                docxs[Enterprice][job].append(norm_f)
                                seen_functions.add(norm_f.upper())
            else:
                error_files[str(element)] = "Not a .docx file"
        elif element.is_dir():
            print(f"@ Explorando subcarpeta: {str(element)}")
            list_files(element, docxs)
    return docxs 

def main():
    directory = Path("./Files")
    dir_json = "./Results/docx_files.json"
    list_docxs = list_files(directory)
    with open(dir_errors, 'w', encoding='utf-8') as f:
        json.dump(error_files, f, ensure_ascii=False, indent=2)
    print(f"Procesados {len(list_docxs)} archivos.")
    with open(dir_json, 'w', encoding='utf-8') as f:
        json.dump(list_docxs, f, ensure_ascii=False, indent=2)
    conjunction_dbs(list_docxs)
    validate_no_overlap()
    generate_sql_script(list_docxs)
    corregir_duplicados_orden()
    generar_reporte_items()
    
if __name__ == "__main__":
    main()

